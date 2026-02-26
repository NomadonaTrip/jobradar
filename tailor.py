#!/usr/bin/env python3
"""
Auto-Tailor — Phase 2 of the autonomous resume-tailoring pipeline.

For each new JD in the JDs/ folder, generates:
  1. Tailored resume (Markdown + DOCX)
  2. Cover letter (Markdown)
  3. Match analysis report (Markdown)

Uses the `claude` CLI (your existing subscription) as the AI backend.

Usage:
    python tailor.py                    # process all untailored JDs
    python tailor.py --limit 3          # process only 3 JDs
    python tailor.py --jd "CGI*.md"     # process specific JD(s) by glob
    python tailor.py --dry-run          # preview what would be processed
"""

import argparse
import fnmatch
import json
import os
import re
import subprocess
import sys
import time
from datetime import datetime, timezone
from pathlib import Path

import yaml

# ---------------------------------------------------------------------------
# AI Cliche Cleanup — Post-processing layer
# ---------------------------------------------------------------------------

# Each stem maps morphological forms to cleaner replacements.
# The sanitize function handles case preservation automatically.
BANNED_WORD_MAP = {
    "leverage":     {"leveraged": "used", "leveraging": "using", "leverages": "uses", "leverage": "use"},
    "utilize":      {"utilized": "used", "utilizing": "using", "utilizes": "uses", "utilize": "use"},
    "spearhead":    {"spearheaded": "led", "spearheading": "leading", "spearheads": "leads", "spearhead": "lead"},
    "facilitate":   {"facilitated": "coordinated", "facilitating": "coordinating", "facilitates": "coordinates", "facilitate": "coordinate"},
    "foster":       {"fostered": "built", "fostering": "building", "fosters": "builds", "foster": "build"},
    "cultivate":    {"cultivated": "developed", "cultivating": "developing", "cultivates": "develops", "cultivate": "develop"},
    "harness":      {"harnessed": "used", "harnessing": "using", "harnesses": "uses", "harness": "use"},
    "champion":     {"championed": "led", "championing": "leading", "champions": "leads", "champion": "lead"},
    "seamless":     {"seamlessly": "smoothly", "seamless": "smooth"},
    "robust":       {"robust": "strong"},
    "holistic":     {"holistically": "comprehensively", "holistic": "comprehensive"},
    "cutting-edge": {"cutting-edge": "modern"},
    "best-in-class":{"best-in-class": "leading"},
    "synergy":      {"synergies": "collaboration", "synergy": "collaboration"},
    "paradigm":     {"paradigms": "approaches", "paradigm": "approach"},
}

# Phrases to strip entirely or replace. Format: (regex_pattern, replacement).
BANNED_PHRASES = [
    (r"I am writing to express my (?:strong |keen )?interest in", "I am applying for"),
    (r"resonates deeply with", "matches"),
    (r"aligns perfectly with", "matches"),
    (r"I am (?:excited|thrilled|eager) (?:about|by) the opportunity to", "I look forward to"),
    (r"I would welcome the (?:opportunity|chance) to discuss", "I am happy to discuss"),
    (r"(?:a |the )?culture of (\w+)", r"\1"),
    (r"in today'?s (?:\w+ )*?landscape", ""),
    (r"at the forefront of", ""),
    (r"(?:deep|strong) (?:understanding|passion) (?:of|for)", "experience with"),
    (r"actionable (insights?|recommendations?|deliverables?)", r"clear \1"),
    (r"well-?versed in", "experienced with"),
    (r"uniquely positioned to", "prepared to"),
    (r"well-?positioned to", "prepared to"),
]


def _preserve_case(original: str, replacement: str) -> str:
    """Apply the case pattern of *original* to *replacement*."""
    if not original or not replacement:
        return replacement
    if original.isupper():
        return replacement.upper()
    if original[0].isupper() and (len(original) == 1 or original[1:].islower()):
        return replacement[0].upper() + replacement[1:]
    return replacement


def sanitize_ai_output(text: str, context: str = "resume") -> str:
    """Clean AI-generated text of cliche language patterns.

    context:
      "resume"       - full cleanup (word map + phrases + em-dashes)
      "cover_letter" - full cleanup
      "report"       - light cleanup (em-dashes + whitespace only)
    """
    if not text:
        return text

    # --- Em-dash cleanup (all contexts) ---
    text = re.sub(r"^—\s*", "- ", text, flags=re.MULTILINE)  # list-item starter
    text = text.replace(" — ", " - ")   # inline parenthetical with spaces
    text = text.replace("— ", " - ")    # leading-space variant
    text = text.replace(" —", " -")     # trailing variant
    text = text.replace("—", " - ")     # bare em-dash fallback

    # --- Word map replacements (resume and cover_letter only) ---
    if context in ("resume", "cover_letter"):
        for _stem, forms in BANNED_WORD_MAP.items():
            for original_form, replacement_form in forms.items():
                if "-" in original_form:
                    pattern = re.compile(
                        r"(?<![a-zA-Z])" + re.escape(original_form) + r"(?![a-zA-Z])",
                        re.IGNORECASE,
                    )
                else:
                    pattern = re.compile(
                        r"\b" + re.escape(original_form) + r"\b",
                        re.IGNORECASE,
                    )

                def _make_replacer(repl):
                    def replacer(match):
                        return _preserve_case(match.group(0), repl)
                    return replacer

                text = pattern.sub(_make_replacer(replacement_form), text)

    # --- Phrase cleanup (resume and cover_letter only) ---
    if context in ("resume", "cover_letter"):
        for phrase_pattern, phrase_replacement in BANNED_PHRASES:
            text = re.sub(phrase_pattern, phrase_replacement, text, flags=re.IGNORECASE)

    # --- Whitespace normalization (all contexts) ---
    text = re.sub(r"([^\n]) {2,}", r"\1 ", text)        # collapse multi-spaces
    text = re.sub(r" +$", "", text, flags=re.MULTILINE)  # trim trailing spaces
    text = re.sub(r"\n{4,}", "\n\n\n", text)             # collapse excess blank lines

    return text


# ---------------------------------------------------------------------------
# Paths & Config
# ---------------------------------------------------------------------------
ROOT = Path(os.environ["PIPELINE_WORKDIR"]) if "PIPELINE_WORKDIR" in os.environ else Path(__file__).resolve().parent
CONFIG_PATH = ROOT / "config.yaml"
TAILOR_STATE_FILE = ROOT / "tailor_state.json"


def load_config() -> dict:
    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def load_tailor_state() -> dict:
    if TAILOR_STATE_FILE.exists():
        with open(TAILOR_STATE_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        data.setdefault("tailored", {})
        return data
    return {"tailored": {}}


def save_tailor_state(state: dict):
    with open(TAILOR_STATE_FILE, "w", encoding="utf-8") as f:
        json.dump(state, f, indent=2, ensure_ascii=False)


# ---------------------------------------------------------------------------
# DOCX generation from Markdown
# ---------------------------------------------------------------------------
def md_to_docx(md_path: Path, docx_path: Path):
    """Convert a Markdown resume to a simple DOCX file using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print(f"    [DOCX] python-docx not installed, skipping DOCX generation.")
        return

    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    md_text = md_path.read_text(encoding="utf-8")

    for line in md_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        # H1 — Name
        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # H2 — Section headers
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        # H3 — Job titles
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        # Bullet points
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        # Bold lines (like dates, certifications)
        elif stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True
            run.font.size = Pt(10)
        # Horizontal rules
        elif stripped == "---":
            continue
        # Regular text
        else:
            # Strip markdown bold/italic for plain text
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
            clean = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", clean)
            if clean:
                p = doc.add_paragraph(clean)
                p.paragraph_format.space_after = Pt(2)

    doc.save(str(docx_path))


# ---------------------------------------------------------------------------
# Claude CLI integration
# ---------------------------------------------------------------------------
def _clean_env() -> dict:
    """Return a copy of os.environ without CLAUDECODE so nested sessions work."""
    env = os.environ.copy()
    env.pop("CLAUDECODE", None)
    return env


def call_claude(prompt: str, max_retries: int = 2) -> str:
    """Call the claude CLI in non-interactive mode and return the response."""
    env = _clean_env()
    for attempt in range(max_retries + 1):
        try:
            result = subprocess.run(
                [
                    "claude", "-p", prompt,
                    "--model", "sonnet",
                    "--no-session-persistence",
                    "--output-format", "text",
                ],
                capture_output=True,
                text=True,
                timeout=300,  # 5 minute timeout per call
                cwd=str(ROOT),
                env=env,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
            elif result.returncode != 0:
                err = result.stderr.strip() or result.stdout.strip()
                print(f"    [Claude] Attempt {attempt+1} failed (exit {result.returncode}): {err[:200]}")
                if attempt < max_retries:
                    time.sleep(5)
                    continue
                return ""
        except subprocess.TimeoutExpired:
            print(f"    [Claude] Attempt {attempt+1} timed out.")
            if attempt < max_retries:
                time.sleep(5)
                continue
            return ""
    return ""


# ---------------------------------------------------------------------------
# Autonomous skill integration
# ---------------------------------------------------------------------------
AUTONOMOUS_SKILL_PATH = Path.home() / ".claude" / "skills" / "resume-tailoring-autonomous" / "SKILL.md"


def load_autonomous_skill() -> str:
    """Load the autonomous skill system prompt. Exits if missing."""
    if not AUTONOMOUS_SKILL_PATH.exists():
        print(f"CRITICAL: Autonomous skill not found at {AUTONOMOUS_SKILL_PATH}")
        print(f"Cannot proceed with tailoring. The skill file is required.")
        sys.exit(1)
    content = AUTONOMOUS_SKILL_PATH.read_text(encoding="utf-8")
    print(f"  Autonomous skill loaded: {AUTONOMOUS_SKILL_PATH.name} ({len(content)} chars)")
    return content


def load_resume_library(config: dict) -> str:
    """Load all resume files from the library directory into a delimited string.

    Falls back to base_resume.md if resumes/ directory doesn't exist (backward compat).
    Returns empty string if no resume content found.
    """
    library_dir = ROOT / config["candidate"].get("resume_library", "resumes")
    parts = []

    if library_dir.exists() and library_dir.is_dir():
        for md_file in sorted(library_dir.glob("*.md")):
            content = md_file.read_text(encoding="utf-8")
            if content.strip():
                parts.append(f"=== RESUME FILE: {md_file.name} ===\n{content}\n=== END FILE ===")
        if parts:
            print(f"  Resume library: {len(parts)} file(s) from {library_dir.name}/")
            return "\n\n".join(parts)

    # Fallback: use base_resume.md
    base_path = ROOT / config["candidate"].get("base_resume", "base_resume.md")
    if base_path.exists():
        content = base_path.read_text(encoding="utf-8")
        if content.strip():
            print(f"  Resume library: fallback to {base_path.name}")
            return f"=== RESUME FILE: {base_path.name} ===\n{content}\n=== END FILE ==="

    return ""


def load_cover_letter_library(config: dict) -> str:
    """Load all cover letter files and concatenate for voice reference.

    Falls back to base_cover_letter.md if cover_letters/ doesn't exist.
    Returns empty string if no cover letter content found.
    """
    cl_dir = ROOT / config["candidate"].get("cover_letter_library", "cover_letters")
    parts = []

    if cl_dir.exists() and cl_dir.is_dir():
        for md_file in sorted(cl_dir.glob("*.md")):
            content = md_file.read_text(encoding="utf-8")
            if content.strip():
                parts.append(content)
        if parts:
            print(f"  Cover letter library: {len(parts)} file(s) from {cl_dir.name}/")
            return "\n\n---\n\n".join(parts)

    # Fallback: use base_cover_letter.md
    cl_path = ROOT / config["candidate"].get("base_cover_letter", "base_cover_letter.md")
    if cl_path.exists():
        content = cl_path.read_text(encoding="utf-8")
        if content.strip():
            print(f"  Cover letter library: fallback to {cl_path.name}")
            return content

    return ""


def _extract_text_from_json_output(raw_json: str) -> str:
    """Extract the best text content from claude JSON output.

    When --output-format json is used with tool-enabled sessions, the response
    contains the full conversation. The delimited resume/report may be in an
    intermediate message (before the model's final summary). We scan all
    assistant messages for the delimiters and return the first match.
    Falls back to the final assistant message text if no delimiters found.
    """
    try:
        data = json.loads(raw_json)
    except json.JSONDecodeError:
        return raw_json  # Not JSON, return as-is

    # Handle the JSON response structure from claude CLI
    # The result field contains the final text, but messages contains all turns
    messages = data.get("messages", [])
    result_text = data.get("result", "")

    # Scan all assistant messages for delimiter markers (prefer the one with markers)
    for msg in messages:
        if msg.get("role") != "assistant":
            continue
        content = msg.get("content", "")
        # content may be a string or a list of content blocks
        if isinstance(content, list):
            text_parts = []
            for block in content:
                if isinstance(block, dict) and block.get("type") == "text":
                    text_parts.append(block["text"])
            content = "\n".join(text_parts)
        if "<<RESUME_START>>" in content:
            return content

    # No delimiters found in any message — fall back to result field
    if result_text:
        return result_text

    # Last resort: concatenate all assistant text blocks
    all_text = []
    for msg in messages:
        if msg.get("role") != "assistant":
            continue
        content = msg.get("content", "")
        if isinstance(content, list):
            for block in content:
                if isinstance(block, dict) and block.get("type") == "text":
                    all_text.append(block["text"])
        elif isinstance(content, str):
            all_text.append(content)
    return "\n".join(all_text) if all_text else raw_json


def call_claude_with_skill(prompt: str, skill_prompt: str, max_retries: int = 2) -> str:
    """Call claude CLI with the autonomous skill as a system prompt append.

    Uses --output-format json to capture the full conversation including
    intermediate messages from tool use (WebSearch). This ensures we can
    extract the delimited resume/report even if the model's final message
    is a summary.
    """
    env = _clean_env()
    for attempt in range(max_retries + 1):
        try:
            result = subprocess.run(
                [
                    "claude", "-p", prompt,
                    "--model", "sonnet",
                    "--no-session-persistence",
                    "--output-format", "json",
                    "--append-system-prompt", skill_prompt,
                    "--allowedTools", "WebSearch",
                    "--max-budget-usd", "3.00",
                ],
                capture_output=True,
                text=True,
                timeout=600,  # 10 minute timeout — skill methodology is more complex
                cwd=str(ROOT),
                env=env,
            )
            if result.returncode == 0 and result.stdout.strip():
                return _extract_text_from_json_output(result.stdout.strip())
            elif result.returncode != 0:
                err = result.stderr.strip() or result.stdout.strip()
                print(f"    [Skill] Attempt {attempt+1} failed (exit {result.returncode}): {err[:200]}")
                if attempt < max_retries:
                    time.sleep(5)
                    continue
                return ""
        except subprocess.TimeoutExpired:
            print(f"    [Skill] Attempt {attempt+1} timed out (10 min).")
            if attempt < max_retries:
                time.sleep(5)
                continue
            return ""
    return ""


def extract_between(text: str, start_marker: str, end_marker: str) -> str:
    """Extract text between markers, returning empty string if not found."""
    start_idx = text.find(start_marker)
    end_idx = text.find(end_marker)
    if start_idx == -1 or end_idx == -1 or end_idx <= start_idx:
        return ""
    return text[start_idx + len(start_marker):end_idx].strip()


def _extract_jd_metadata(jd_text: str) -> dict:
    """Extract structured metadata from a JD markdown file.

    Parses the H1 heading (# Role - Company) and the metadata table
    (| **Field** | Value |) to return {company, role, location}.
    """
    meta = {"company": "", "role": "", "location": ""}

    for line in jd_text.split("\n"):
        stripped = line.strip()

        # H1 heading: "# Role — Company"
        if stripped.startswith("# ") and not meta["role"]:
            heading = stripped[2:]
            if " — " in heading:
                meta["role"], meta["company"] = heading.split(" — ", 1)
            elif " - " in heading:
                meta["role"], meta["company"] = heading.split(" - ", 1)
            else:
                meta["role"] = heading

        # Metadata table rows
        if "**Company**" in stripped and "|" in stripped:
            parts = stripped.split("|")
            if len(parts) >= 3:
                meta["company"] = parts[2].strip().strip("*")
        if "**Location**" in stripped and "|" in stripped:
            parts = stripped.split("|")
            if len(parts) >= 3:
                meta["location"] = parts[2].strip().strip("*")

    meta = {k: v.strip() for k, v in meta.items()}
    return meta


def build_skill_resume_prompt(resume_library: str, jd_text: str, candidate_name: str,
                               discovery_notes: str = "") -> str:
    """Build the prompt for skill-based resume + report generation."""
    discovery_section = ""
    if discovery_notes:
        discovery_section = f"""
DISCOVERY NOTES (candidate's self-reported achievements, metrics, and context from onboarding):
---
{discovery_notes}
---

"""

    # Extract structured metadata so the skill can build targeted research queries
    meta = _extract_jd_metadata(jd_text)
    metadata_section = ""
    if meta.get("company") or meta.get("role"):
        metadata_section = f"""STRUCTURED JOB METADATA (use these for targeted WebSearch queries):
- Company: {meta.get('company', 'Unknown')}
- Role: {meta.get('role', 'Unknown')}
- Location: {meta.get('location', 'Not specified')}

"""

    return f"""TASK: Tailor a resume for {candidate_name} using the autonomous pipeline skill methodology.

MODE: EXPRESS / HEADLESS - Execute all phases automatically. No checkpoints. No interactive questions.

RESUME LIBRARY:
{resume_library}

{discovery_section}{metadata_section}TARGET JOB DESCRIPTION:
---
{jd_text}
---

Execute the full skill workflow (Library Parse -> JD Analysis + Company Research -> Template -> Content Match -> Generate).

WRITING QUALITY REMINDER: Follow the anti-cliche rules in the skill prompt strictly. No em-dashes, no "leveraged", no "spearheaded", no "facilitated", no "seamless", no "actionable insights", no "foster a culture of". Use plain, strong resume verbs (led, built, delivered, managed, reduced, improved, designed, implemented). Vary sentence openings across bullets.

CRITICAL OUTPUT FORMAT - Your ENTIRE response must use these exact delimiters and contain NOTHING else:

<<RESUME_START>>
(the full tailored resume in Markdown goes here)
<<RESUME_END>>

<<REPORT_START>>
(the full match analysis report in Markdown goes here)
<<REPORT_END>>

Do NOT output any commentary, preamble, summary, or text outside these delimiters. The output is parsed programmatically.

Begin now."""


# ---------------------------------------------------------------------------
# Prompts (legacy - kept for cover letter which is separate from the skill)
# ---------------------------------------------------------------------------
def build_resume_prompt(base_resume: str, jd_text: str, candidate_name: str) -> str:
    return f"""You are an expert resume writer specializing in ATS optimization and strategic keyword matching for senior Scrum Master and Product Manager roles.

TASK: Tailor the candidate's base resume to match the target job description. Output ONLY the tailored resume in Markdown format - no commentary, no explanations, no preamble.

RULES:
1. TRUTHFULNESS IS PARAMOUNT - never fabricate experience, metrics, or skills. Only reframe existing experience using JD-aligned language.
2. Mirror the JD's exact terminology wherever truthful (e.g., if JD says "high-performing teams" use that phrase, if JD says "processing components" use that phrase).
3. Reorder bullet points so the most JD-relevant ones come first under each role.
4. Adjust the Professional Summary to directly address the target role's key requirements.
5. Keep all quantified metrics (78%, 30%, $15M, etc.) - they are powerful differentiators.
6. Maintain the same resume structure: Name -> Professional Summary -> Certifications -> Key Skills -> Professional Experience -> Education.
7. In Key Skills, add any JD-mentioned skills the candidate genuinely has but hasn't listed.
8. Keep to 1-2 pages. Be concise but impactful.
9. Do NOT include the company name or job title in the resume header - it should work as a general submission.

BASE RESUME:
---
{base_resume}
---

TARGET JOB DESCRIPTION:
---
{jd_text}
---

Output the tailored resume in Markdown now:"""


def build_cover_letter_prompt(base_resume: str, jd_text: str, tailored_resume: str, candidate_name: str, company: str, role: str, cover_letter_voice: str = "") -> str:
    voice_section = ""
    if cover_letter_voice:
        voice_section = f"""
VOICE REFERENCE - Match this candidate's authentic writing style, tone, and personality.
Study the sentence structure, warmth level, formality, and storytelling approach in the sample below.
Mirror how they open letters, transition between ideas, and express enthusiasm. Do NOT copy content - only match the voice.
---
{cover_letter_voice}
---

"""
    return f"""You are writing a cover letter for {candidate_name} applying to the {role} position at {company}.
{voice_section}VOICE & TONE:
- {"Match the candidate's voice from the reference above" if cover_letter_voice else "Professional but warm and human"} - not robotic or generic
- Show genuine interest in the company's mission and culture
- Weave in specific achievements with metrics
- Keep it to 3-4 paragraphs, under 400 words
- Address "Dear Hiring Manager" unless a specific name is in the JD
- End with a confident but not arrogant call to action

WRITING QUALITY:
- Do NOT use em-dashes. Use commas, colons, or periods instead.
- Do NOT open with "I am writing to express my strong interest" or any variation. Start with something specific about the company or a direct statement of fit.
- Do NOT use "resonates deeply", "aligns perfectly", "I am excited about the opportunity", or "I would welcome the opportunity to discuss".
- Banned words: leverage, utilize, spearhead, facilitate, foster, cultivate, harness, holistic, robust, cutting-edge, seamless, synergy, paradigm, actionable.
- Use plain, confident language: "use" not "utilize", "led" not "spearheaded", "strong" not "robust".
- Each paragraph should have a distinct purpose. Avoid restating the same point in different buzzwords.
- Close with a specific, natural call to action, not a formula.

STRUCTURE:
1. "Dear Hiring Manager," (or specific name if in JD)
2. Opening paragraph — Hook connecting candidate's passion/experience to the company's mission. Reference something specific about the company.
3. Body (1-2 paragraphs) — 3-4 strongest achievements that directly map to JD requirements. Use specific numbers.
4. Closing paragraph — Express enthusiasm, mention willingness to discuss further.
5. "Sincerely," followed by candidate name.

IMPORTANT: Do NOT include the candidate's address, contact info, or resume header. Start directly with the greeting.

CANDIDATE'S RESUME:
---
{tailored_resume}
---

JOB DESCRIPTION:
---
{jd_text}
---

Output ONLY the cover letter text in Markdown format — no commentary:"""


def build_report_prompt(base_resume: str, jd_text: str, tailored_resume: str, company: str, role: str) -> str:
    return f"""You are a career strategist analyzing the fit between a candidate and a target role.

Generate a Match Analysis Report in Markdown with these exact sections:

## Target Role Summary
Company, position, location, salary (if mentioned), key details from JD.

## Content Mapping Summary
Table showing: total bullets, direct matches (%), transferable matches (%), adjacent matches (%), gaps (%).
Overall JD Coverage percentage.

## Key Reframings Applied
Table: Original phrasing → Reframed phrasing → Reason for change.
Note that all reframings are truthful.

## Gap Analysis
### Identified Gaps
Table: Gap | Severity (Low/Medium/High) | Mitigation strategy.
### Strengths vs. JD
Table: JD Requirement | Resume Evidence | Confidence (%).

## Key Differentiators
Numbered list of 4-6 things that make this candidate stand out for THIS specific role.

## Recommendations for Interview Prep
### Stories to Prepare
5-7 specific STAR-format stories to prepare, drawn from the resume.
### Questions to Expect
5-6 likely interview questions specific to this role.
### Gaps to Address Proactively
How to spin each gap positively in conversation.

BASE RESUME (before tailoring):
---
{base_resume}
---

TAILORED RESUME (after tailoring):
---
{tailored_resume}
---

JOB DESCRIPTION:
---
{jd_text}
---

Output the full report in Markdown now:"""


# ---------------------------------------------------------------------------
# Process a single JD
# ---------------------------------------------------------------------------
def process_jd(jd_path: Path, base_resume: str, config: dict, dry_run: bool = False,
               cover_letter_voice: str = "", resume_library: str = "",
               discovery_notes: str = "", skill_prompt: str = "") -> bool:
    """Process a single JD file: generate tailored resume + report (skill-based), then cover letter."""
    jd_text = jd_path.read_text(encoding="utf-8")
    candidate = config["candidate"]
    candidate_name = candidate["name"]

    # Extract company and role from JD content or filename
    company = "Unknown"
    role = "Unknown"

    # Strategy 1: Parse the H1 heading (JSearch-fetched JDs use "# Role — Company")
    for line in jd_text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# "):
            heading = stripped[2:]
            if " — " in heading:
                role, company = heading.split(" — ", 1)
            elif " - " in heading:
                role, company = heading.split(" - ", 1)
            else:
                role = heading
            break

    # Strategy 2: Look for Company in the metadata table
    if company == "Unknown":
        for line in jd_text.split("\n"):
            if "**Company**" in line and "|" in line:
                company = line.split("|")[2].strip().strip("*")
                break

    # Strategy 3: Fall back to filename (format: "Company_Role.md")
    if company == "Unknown" or role == "Unknown":
        stem = jd_path.stem
        parts = stem.split("_", 1)
        if company == "Unknown":
            company = parts[0].replace("-", " ") if parts else "Unknown"
        if role == "Unknown" and len(parts) > 1:
            role = parts[1].replace("_", " ")

    company = company.strip()
    role = role.strip()

    print(f"\n  Processing: {role} @ {company}")
    print(f"  JD file: {jd_path.name}")

    if dry_run:
        print(f"  [DRY RUN] Would generate: resume + report (skill-based), cover letter")
        return True

    # Create output directory
    safe_name = re.sub(r'[<>:"/\\|?*]', "", f"{company}_{role}")
    safe_name = re.sub(r"\s+", "_", safe_name)[:100]
    output_dir = ROOT / "output" / safe_name
    output_dir.mkdir(parents=True, exist_ok=True)

    # Save JD copy to output
    (output_dir / "jd.md").write_text(jd_text, encoding="utf-8")

    # Step 1: Skill-based tailored resume + report (single call)
    print(f"  [1/2] Generating tailored resume + report (skill-based)...")
    prompt = build_skill_resume_prompt(resume_library, jd_text, candidate_name, discovery_notes)
    raw_output = call_claude_with_skill(prompt, skill_prompt)

    if not raw_output:
        print(f"  ERROR: Failed to generate tailored resume. Skipping.")
        return False

    # Parse delimited output
    tailored_resume = extract_between(raw_output, "<<RESUME_START>>", "<<RESUME_END>>")
    report = extract_between(raw_output, "<<REPORT_START>>", "<<REPORT_END>>")

    if not tailored_resume:
        # Fallback: treat entire output as resume if markers not found
        print(f"    WARNING: Resume markers not found in output, using full response as resume.")
        tailored_resume = raw_output

    tailored_resume = sanitize_ai_output(tailored_resume, context="resume")
    resume_md_path = output_dir / f"{candidate_name.replace(' ', '_')}_{safe_name}_Resume.md"
    resume_md_path.write_text(tailored_resume, encoding="utf-8")
    print(f"    Saved: {resume_md_path.name}")

    # Generate DOCX
    docx_path = resume_md_path.with_suffix(".docx")
    md_to_docx(resume_md_path, docx_path)
    print(f"    Saved: {docx_path.name}")

    # Save report
    if report:
        report = sanitize_ai_output(report, context="report")
        report_path = output_dir / f"{candidate_name.replace(' ', '_')}_{safe_name}_Report.md"
        report_path.write_text(report, encoding="utf-8")
        print(f"    Saved: {report_path.name}")
    else:
        print(f"    WARNING: Report markers not found in output. Report not generated.")

    # Step 2: Cover letter (separate call, uses all cover letter voice sources)
    if cover_letter_voice:
        print(f"  [2/2] Generating cover letter...")
        cl_prompt = build_cover_letter_prompt(base_resume, jd_text, tailored_resume, candidate_name, company, role, cover_letter_voice)
        cover_letter = call_claude(cl_prompt)
        if cover_letter:
            cover_letter = sanitize_ai_output(cover_letter, context="cover_letter")
            cl_path = output_dir / f"{candidate_name.replace(' ', '_')}_{safe_name}_CoverLetter.md"
            cl_path.write_text(cover_letter, encoding="utf-8")
            print(f"    Saved: {cl_path.name}")
        else:
            print(f"    WARNING: Cover letter generation failed.")
    else:
        print(f"  [2/2] Skipping cover letter — no cover letter voice reference available.")

    return True


# ---------------------------------------------------------------------------
# Main pipeline
# ---------------------------------------------------------------------------
def get_unprocessed_jds(jd_dir: Path, state: dict, jd_glob: str | None = None) -> list[Path]:
    """Find JD files that haven't been tailored yet."""
    all_jds = sorted(jd_dir.glob("*.md"))

    if jd_glob:
        all_jds = [p for p in all_jds if fnmatch.fnmatch(p.name, jd_glob)]

    unprocessed = []
    for jd in all_jds:
        if jd.name not in state.get("tailored", {}):
            unprocessed.append(jd)

    return unprocessed


def run(dry_run: bool = False, limit: int | None = None, jd_glob: str | None = None):
    print("=" * 60)
    print(f"  Auto-Tailor — {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    print("=" * 60)

    config = load_config()
    state = load_tailor_state()

    # Load autonomous skill (required — exits if missing)
    skill_prompt = load_autonomous_skill()

    # Load resume library (all sources from resumes/ directory)
    resume_library = load_resume_library(config)
    if not resume_library:
        print(f"ERROR: No resume content found. Add resumes to {ROOT / 'resumes'}/")
        print(f"Resume tailoring requires at least one real resume source.")
        sys.exit(1)

    # Load base resume for backward compat (used in cover letter prompt)
    base_resume_path = ROOT / config["candidate"].get("base_resume", "base_resume.md")
    base_resume = ""
    if base_resume_path.exists():
        base_resume = base_resume_path.read_text(encoding="utf-8")

    # Load cover letter voice from all sources
    cover_letter_voice = load_cover_letter_library(config)
    if not cover_letter_voice:
        print(f"  NOTE: No cover letter voice reference found. Cover letters will be skipped.")

    # Load discovery notes (optional supplement)
    discovery_notes = ""
    discovery_path = ROOT / "discovery_notes.md"
    if discovery_path.exists():
        discovery_notes = discovery_path.read_text(encoding="utf-8")
        print(f"  Discovery notes: {discovery_path.name} ({len(discovery_notes)} chars)")

    # Find unprocessed JDs
    jd_dir = ROOT / config["output"]["jd_directory"]
    unprocessed = get_unprocessed_jds(jd_dir, state, jd_glob)

    if limit:
        unprocessed = unprocessed[:limit]

    print(f"\n  JDs to process: {len(unprocessed)}")
    if not unprocessed:
        print("\n  No new JDs to tailor. Run fetcher.py first or use --jd to specify a pattern.")
        return

    print(f"  Output directory: {ROOT / 'output'}/")
    print()

    # Process each JD
    success = 0
    failed = 0
    for i, jd_path in enumerate(unprocessed, 1):
        print(f"  [{i}/{len(unprocessed)}]", end="")
        ok = process_jd(jd_path, base_resume, config, dry_run=dry_run,
                        cover_letter_voice=cover_letter_voice,
                        resume_library=resume_library,
                        discovery_notes=discovery_notes,
                        skill_prompt=skill_prompt)

        if ok and not dry_run:
            state["tailored"][jd_path.name] = {
                "processed_at": datetime.now(timezone.utc).isoformat(),
                "output_dir": str(ROOT / "output"),
            }
            save_tailor_state(state)
            success += 1
        elif ok and dry_run:
            success += 1
        else:
            failed += 1

    # Summary
    print("\n" + "=" * 60)
    print(f"  TAILORING COMPLETE")
    print("=" * 60)
    print(f"  Processed: {success + failed}")
    print(f"  Success:   {success}")
    print(f"  Failed:    {failed}")
    if not dry_run:
        print(f"  Output at: {ROOT / 'output'}/")
    print()


# ---------------------------------------------------------------------------
# Cover letter backfill
# ---------------------------------------------------------------------------
def backfill_cover_letters(dry_run: bool = False, limit: int | None = None):
    """Generate cover letters for output packages that are missing them."""
    print("=" * 60)
    print(f"  Cover Letter Backfill — {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    print("=" * 60)

    config = load_config()
    candidate = config["candidate"]
    candidate_name = candidate["name"]

    # Load base resume (used in cover letter prompt)
    base_resume_path = ROOT / candidate.get("base_resume", "base_resume.md")
    base_resume = ""
    if base_resume_path.exists():
        base_resume = base_resume_path.read_text(encoding="utf-8")

    # Load cover letter voice — required for backfill
    cover_letter_voice = load_cover_letter_library(config)
    if not cover_letter_voice:
        print(f"\n  ERROR: No cover letter voice reference found.")
        print(f"  Backfill requires base_cover_letter.md or files in cover_letters/.")
        sys.exit(1)

    # Scan output dirs for packages missing a cover letter
    output_root = ROOT / "output"
    if not output_root.exists():
        print(f"\n  No output directory found at {output_root}")
        return

    missing = []
    for pkg_dir in sorted(output_root.iterdir()):
        if not pkg_dir.is_dir():
            continue
        has_resume = list(pkg_dir.glob("*_Resume.md"))
        has_jd = (pkg_dir / "jd.md").exists()
        has_cover = list(pkg_dir.glob("*CoverLetter.md"))
        if has_resume and has_jd and not has_cover:
            missing.append(pkg_dir)

    if limit:
        missing = missing[:limit]

    print(f"\n  Packages missing cover letters: {len(missing)}")
    if not missing:
        print("  Nothing to backfill.")
        return

    print(f"  Output directory: {output_root}/")
    print()

    success = 0
    failed = 0
    for i, pkg_dir in enumerate(missing, 1):
        safe_name = pkg_dir.name
        jd_text = (pkg_dir / "jd.md").read_text(encoding="utf-8")
        resume_path = list(pkg_dir.glob("*_Resume.md"))[0]
        tailored_resume = resume_path.read_text(encoding="utf-8")

        # Extract company and role from JD (same logic as process_jd)
        company = "Unknown"
        role = "Unknown"
        for line in jd_text.split("\n"):
            stripped = line.strip()
            if stripped.startswith("# "):
                heading = stripped[2:]
                if " — " in heading:
                    role, company = heading.split(" — ", 1)
                elif " - " in heading:
                    role, company = heading.split(" - ", 1)
                else:
                    role = heading
                break
        if company == "Unknown":
            for line in jd_text.split("\n"):
                if "**Company**" in line and "|" in line:
                    company = line.split("|")[2].strip().strip("*")
                    break
        company = company.strip()
        role = role.strip()

        print(f"  [{i}/{len(missing)}] {role} @ {company}")

        if dry_run:
            print(f"    [DRY RUN] Would generate cover letter in {safe_name}/")
            success += 1
            continue

        cl_prompt = build_cover_letter_prompt(
            base_resume, jd_text, tailored_resume,
            candidate_name, company, role, cover_letter_voice
        )
        cover_letter = call_claude(cl_prompt)
        if cover_letter:
            cover_letter = sanitize_ai_output(cover_letter, context="cover_letter")
            cl_path = pkg_dir / f"{candidate_name.replace(' ', '_')}_{safe_name}_CoverLetter.md"
            cl_path.write_text(cover_letter, encoding="utf-8")
            print(f"    Saved: {cl_path.name}")
            success += 1
        else:
            print(f"    WARNING: Cover letter generation failed.")
            failed += 1

    print("\n" + "=" * 60)
    print(f"  BACKFILL COMPLETE")
    print("=" * 60)
    print(f"  Processed: {success + failed}")
    print(f"  Success:   {success}")
    print(f"  Failed:    {failed}")
    print()


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Auto-tailor resumes to fetched JDs")
    parser.add_argument("--dry-run", action="store_true", help="Preview without generating")
    parser.add_argument("--limit", type=int, help="Max number of JDs to process")
    parser.add_argument("--jd", type=str, help="Glob pattern to match specific JD files (e.g. 'CGI*.md')")
    parser.add_argument("--cover-letters-only", action="store_true",
                        help="Backfill cover letters for existing packages that are missing them")
    args = parser.parse_args()

    if args.cover_letters_only:
        backfill_cover_letters(dry_run=args.dry_run, limit=args.limit)
    else:
        run(dry_run=args.dry_run, limit=args.limit, jd_glob=args.jd)
