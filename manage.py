#!/usr/bin/env python3
"""
Customer Management CLI — concierge operations for multiple customers.

Usage:
    python manage.py import onboarding_Tayo_Fasunon.json   # import from onboarding form
    python manage.py list                                   # list all customers
    python manage.py status john-doe                        # show customer status
    python manage.py run john-doe                           # run pipeline for one customer
    python manage.py run-all                                # run pipeline for all customers
    python manage.py run-all --fetch-only                   # just fetch for everyone
"""

import argparse
import base64
import json
import io
import os
import re
import shutil
import subprocess
import sys
from datetime import datetime, timedelta, timezone
from pathlib import Path

import yaml

ROOT = Path(__file__).resolve().parent
CUSTOMERS_DIR = ROOT / "customers"
PYTHON = str(ROOT / ".venv" / "bin" / "python")
TEMPLATE_CONFIG = ROOT / "config.yaml"


def slugify(name: str) -> str:
    """Convert a name to a URL-safe slug."""
    slug = name.lower().strip()
    slug = re.sub(r"[^a-z0-9]+", "-", slug)
    return slug.strip("-")


# ---------------------------------------------------------------------------
# Import from onboarding JSON
# ---------------------------------------------------------------------------
def do_import(data: dict, force: bool = False) -> Path:
    """Core import logic — create a customer directory from onboarding JSON data.

    Returns the customer directory path. Used by both cmd_import (CLI) and
    auto_import.py (Drive poller).
    """
    name = f"{data['firstName']} {data['lastName']}"
    slug = slugify(name)
    customer_dir = CUSTOMERS_DIR / slug

    if customer_dir.exists() and not force:
        raise FileExistsError(f"Customer '{slug}' already exists. Use force=True to overwrite.")

    customer_dir.mkdir(parents=True, exist_ok=True)

    # Decode and save binary resume file if present
    resume_file_data = data.get("resumeFileData", "")
    resume_file_name = data.get("resumeFileName", "") or data.get("resumeFile", "")
    if resume_file_data and resume_file_name:
        try:
            raw_bytes = base64.b64decode(resume_file_data)
            safe_name = re.sub(r"[^a-zA-Z0-9._-]", "_", resume_file_name)
            (customer_dir / safe_name).write_bytes(raw_bytes)
            print(f"  Saved resume file: {safe_name} ({len(raw_bytes)} bytes)")
        except Exception as e:
            print(f"  WARNING: Could not decode/save binary resume: {e}")

    # Save raw onboarding data (strip base64 to keep JSON small)
    data_to_save = {k: v for k, v in data.items() if k not in ("resumeFileData", "coverLetterFileData")}
    (customer_dir / "onboarding.json").write_text(
        json.dumps(data_to_save, indent=2, ensure_ascii=False), encoding="utf-8"
    )

    # Build resume library — store ALL sources, not just one
    resume_library = build_resume_library(data, customer_dir)
    resumes_dir = customer_dir / "resumes"
    resumes_dir.mkdir(exist_ok=True)
    for filename, content in resume_library.items():
        (resumes_dir / filename).write_text(content, encoding="utf-8")
        print(f"  Saved resume source: resumes/{filename} ({len(content)} chars)")

    if not resume_library:
        print(f"  ALERT: No resume data provided (neither pasted nor uploaded).")
        print(f"  Resume tailoring will not work for this customer until a resume is added.")

    # Keep base_resume.md at root for backward compatibility
    best_resume = resume_library.get("pasted_resume.md") or next(iter(resume_library.values()), "")
    if best_resume:
        (customer_dir / "base_resume.md").write_text(best_resume, encoding="utf-8")

    # Decode and save binary cover letter file if present
    cl_file_data = data.get("coverLetterFileData", "")
    cl_file_name = data.get("coverLetterFileName", "") or data.get("coverLetterFile", "")
    if cl_file_data and cl_file_name:
        try:
            cl_raw_bytes = base64.b64decode(cl_file_data)
            cl_safe_name = re.sub(r"[^a-zA-Z0-9._-]", "_", cl_file_name)
            (customer_dir / cl_safe_name).write_bytes(cl_raw_bytes)
            print(f"  Saved cover letter file: {cl_safe_name} ({len(cl_raw_bytes)} bytes)")
        except Exception as e:
            print(f"  WARNING: Could not decode/save binary cover letter: {e}")

    # Build cover letter library — store ALL sources
    cl_library = build_cover_letter_library(data, customer_dir)
    cl_dir = customer_dir / "cover_letters"
    cl_dir.mkdir(exist_ok=True)
    for filename, content in cl_library.items():
        (cl_dir / filename).write_text(content, encoding="utf-8")
        print(f"  Saved cover letter source: cover_letters/{filename} ({len(content)} chars)")

    if not cl_library:
        print(f"  ALERT: No cover letter data provided (neither pasted nor uploaded).")
        print(f"  Cover letter generation will be skipped for this customer.")

    # Keep base_cover_letter.md at root for backward compatibility
    best_cl = cl_library.get("pasted_cover_letter.md") or next(iter(cl_library.values()), "")
    if best_cl:
        (customer_dir / "base_cover_letter.md").write_text(best_cl, encoding="utf-8")

    # Generate customer config
    config = build_customer_config(data, slug)
    with open(customer_dir / "config.yaml", "w", encoding="utf-8") as f:
        yaml.dump(config, f, default_flow_style=False, allow_unicode=True)

    # Create subdirectories (resumes/ and cover_letters/ already created above)
    (customer_dir / "JDs").mkdir(exist_ok=True)
    (customer_dir / "output").mkdir(exist_ok=True)

    # Initialize state files
    for state_file in ["state.json", "tailor_state.json", "notify_state.json"]:
        state_path = customer_dir / state_file
        if not state_path.exists():
            state_path.write_text("{}", encoding="utf-8")

    # Save discovery notes
    discovery = data.get("discovery", {})
    if discovery and any(discovery.values()):
        notes = build_discovery_notes(data)
        (customer_dir / "discovery_notes.md").write_text(notes, encoding="utf-8")

    return customer_dir


def cmd_import(args):
    """Import a customer from onboarding form JSON."""
    json_path = Path(args.file)
    if not json_path.exists():
        print(f"ERROR: File not found: {json_path}")
        sys.exit(1)

    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    try:
        customer_dir = do_import(data, force=args.force)
    except FileExistsError as e:
        print(f"ERROR: {e}")
        sys.exit(1)

    name = f"{data['firstName']} {data['lastName']}"
    slug = slugify(name)
    print(f"\n  Customer imported: {name}")
    print(f"  Slug: {slug}")
    print(f"  Directory: {customer_dir}")
    print(f"  Files created:")
    for f in sorted(customer_dir.rglob("*")):
        if f.is_file():
            print(f"    {f.relative_to(customer_dir)}")
    print(f"\n  Next: python manage.py run {slug}")


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from a PDF resume using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        print("  WARNING: pdfplumber not installed. Run: pip install pdfplumber")
        return ""
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception as e:
        print(f"  WARNING: PDF text extraction failed: {e}")
        return ""
    return "\n\n".join(text_parts)


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a DOCX resume using python-docx."""
    try:
        from docx import Document
    except ImportError:
        print("  WARNING: python-docx not installed.")
        return ""
    text_parts = []
    try:
        doc = Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)
    except Exception as e:
        print(f"  WARNING: DOCX text extraction failed: {e}")
        return ""
    return "\n".join(text_parts)


def _extract_file_text(file_name: str, file_data_b64: str, customer_dir: Path = None) -> str:
    """Extract text from an uploaded file (PDF/DOCX/TXT). Returns empty string on failure."""
    if not file_name:
        return ""
    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
    file_bytes = None

    if file_data_b64:
        try:
            file_bytes = base64.b64decode(file_data_b64)
        except Exception as e:
            print(f"  WARNING: base64 decode failed for {file_name}: {e}")

    if file_bytes is None and customer_dir is not None:
        safe_name = re.sub(r"[^a-zA-Z0-9._-]", "_", file_name)
        disk_path = customer_dir / safe_name
        if disk_path.exists():
            file_bytes = disk_path.read_bytes()

    if file_bytes is None:
        return ""

    if ext == "pdf":
        return _extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return _extract_text_from_docx(file_bytes)
    elif ext in ("txt", "md"):
        try:
            return file_bytes.decode("utf-8")
        except Exception:
            return ""
    return ""


def build_resume_library(data: dict, customer_dir: Path = None) -> dict[str, str]:
    """Build a resume library from ALL onboarding sources.

    Returns a dict of {filename: content} for each source that has content.
    Returns empty dict if neither pasted text nor uploaded file exists.
    """
    library = {}

    # Source 1: Pasted text
    pasted = data.get("resumeText", "").strip()
    if pasted:
        library["pasted_resume.md"] = pasted

    # Source 2: Uploaded file (extract text)
    file_name = data.get("resumeFileName", "") or data.get("resumeFile", "")
    file_data = data.get("resumeFileData", "")
    extracted = _extract_file_text(file_name, file_data, customer_dir).strip()
    if extracted:
        library["uploaded_resume.md"] = extracted
        print(f"  Extracted {len(extracted)} chars from resume file: {file_name}")

    # Source 3: Discovery supplement (only if we have at least one real resume)
    if library:
        discovery = data.get("discovery", {})
        if discovery and any(v.strip() for v in discovery.values() if isinstance(v, str)):
            supplement = _build_discovery_supplement(data)
            if supplement.strip():
                library["discovery_supplement.md"] = supplement

    return library


def _build_discovery_supplement(data: dict) -> str:
    """Build a discovery supplement from onboarding answers. Not a standalone resume."""
    discovery = data.get("discovery", {})
    name = f"{data.get('firstName', '')} {data.get('lastName', '')}".strip()
    parts = [f"# Discovery Supplement: {name}\n"]

    if discovery.get("teamSize"):
        parts.append(f"## Leadership Scale\n{discovery['teamSize']}\n")
    if discovery.get("budget"):
        parts.append(f"## Largest Budget/Project\n{discovery['budget']}\n")
    if discovery.get("metrics"):
        parts.append(f"## Key Metrics & Improvements\n{discovery['metrics']}\n")
    if discovery.get("certs"):
        parts.append(f"## Certifications\n{discovery['certs']}\n")
    if discovery.get("challenge"):
        parts.append(f"## Notable Challenge Overcome\n{discovery['challenge']}\n")
    if discovery.get("tools"):
        parts.append(f"## Tools & Technologies\n{discovery['tools']}\n")
    if discovery.get("industries"):
        parts.append(f"## Industries\n{discovery['industries']}\n")
    if discovery.get("hidden"):
        parts.append(f"## Underrepresented Experience\n{discovery['hidden']}\n")

    return "\n".join(parts)


def build_cover_letter_library(data: dict, customer_dir: Path = None) -> dict[str, str]:
    """Build a cover letter library from ALL onboarding sources.

    Returns a dict of {filename: content}. Empty dict if neither source exists.
    """
    library = {}

    # Source 1: Pasted text
    pasted = data.get("coverLetterText", "").strip()
    if pasted:
        library["pasted_cover_letter.md"] = pasted

    # Source 2: Uploaded file (extract text)
    file_name = data.get("coverLetterFileName", "") or data.get("coverLetterFile", "")
    file_data = data.get("coverLetterFileData", "")
    extracted = _extract_file_text(file_name, file_data, customer_dir).strip()
    if extracted:
        library["uploaded_cover_letter.md"] = extracted
        print(f"  Extracted {len(extracted)} chars from cover letter file: {file_name}")

    return library


def build_resume_from_onboarding(data: dict, customer_dir: Path = None) -> str:
    """Build a base resume markdown from onboarding data.

    DEPRECATED — use build_resume_library() instead. Kept for backward compatibility.
    Now alerts operator instead of building a skeleton when no resume data exists.
    """
    # Priority 1: pasted text
    resume = data.get("resumeText", "")
    if resume:
        return resume

    # Priority 2: extract text from uploaded binary file
    resume_file_name = data.get("resumeFileName", "") or data.get("resumeFile", "")
    resume_file_data = data.get("resumeFileData", "")
    if resume_file_name:
        ext = resume_file_name.rsplit(".", 1)[-1].lower() if "." in resume_file_name else ""
        file_bytes = None

        if resume_file_data:
            try:
                file_bytes = base64.b64decode(resume_file_data)
            except Exception as e:
                print(f"  WARNING: base64 decode failed: {e}")

        if file_bytes is None and customer_dir is not None:
            safe_name = re.sub(r"[^a-zA-Z0-9._-]", "_", resume_file_name)
            disk_path = customer_dir / safe_name
            if disk_path.exists():
                file_bytes = disk_path.read_bytes()

        if file_bytes is not None:
            extracted = ""
            if ext == "pdf":
                extracted = _extract_text_from_pdf(file_bytes)
            elif ext in ("docx", "doc"):
                extracted = _extract_text_from_docx(file_bytes)
            if extracted.strip():
                print(f"  Extracted {len(extracted)} chars from {resume_file_name}")
                return extracted

    # No resume data available — alert operator instead of building a skeleton
    name = f"{data['firstName']} {data['lastName']}"
    print(f"  ALERT: No resume data found for {name} (neither pasted text nor uploaded file).")
    print(f"  A real resume is required for tailoring. Discovery data alone is insufficient.")
    return ""


def build_cover_letter_from_onboarding(data: dict, customer_dir: Path = None) -> str:
    """Build a base cover letter from onboarding data.

    Priority 1: pasted text, Priority 2: extract from uploaded binary file.
    Returns empty string if neither is available.
    """
    # Priority 1: pasted text
    cover_letter = data.get("coverLetterText", "")
    if cover_letter:
        return cover_letter

    # Priority 2: extract text from uploaded binary file
    cl_file_name = data.get("coverLetterFileName", "") or data.get("coverLetterFile", "")
    cl_file_data = data.get("coverLetterFileData", "")
    if cl_file_name:
        ext = cl_file_name.rsplit(".", 1)[-1].lower() if "." in cl_file_name else ""
        file_bytes = None

        if cl_file_data:
            try:
                file_bytes = base64.b64decode(cl_file_data)
            except Exception as e:
                print(f"  WARNING: cover letter base64 decode failed: {e}")

        if file_bytes is None and customer_dir is not None:
            safe_name = re.sub(r"[^a-zA-Z0-9._-]", "_", cl_file_name)
            disk_path = customer_dir / safe_name
            if disk_path.exists():
                file_bytes = disk_path.read_bytes()

        if file_bytes is not None:
            extracted = ""
            if ext == "pdf":
                extracted = _extract_text_from_pdf(file_bytes)
            elif ext in ("docx", "doc"):
                extracted = _extract_text_from_docx(file_bytes)
            if extracted.strip():
                print(f"  Extracted {len(extracted)} chars from cover letter {cl_file_name}")
                return extracted

    return ""


def _build_relevance_config(data: dict) -> dict:
    """Auto-generate relevance scoring config from onboarding data.

    Creates one focus area per role with decreasing weight (first role = 3,
    second = 2, rest = 1). Keywords are derived from role names and certifications.
    """
    roles = data.get("roles", [])
    if not roles:
        return {}

    # Map common role keywords to expanded keyword sets
    ROLE_KEYWORD_MAP = {
        "security": ["security", "cybersecurity", "cyber security", "infosec", "information security"],
        "risk": ["risk", "risk management", "risk assessment", "risk analysis", "enterprise risk"],
        "grc": ["GRC", "governance", "compliance", "audit", "regulatory", "policy"],
        "analyst": ["analyst", "analysis", "assessment", "evaluation"],
        "architect": ["architect", "architecture", "design", "framework"],
        "scrum": ["scrum", "agile", "sprint", "backlog", "scrum master", "kanban"],
        "product": ["product manager", "product owner", "roadmap", "stakeholder", "product strategy"],
        "project": ["project manager", "project management", "PMO", "PMP", "waterfall"],
        "devops": ["devops", "CI/CD", "deployment", "infrastructure", "automation"],
        "cloud": ["cloud", "AWS", "Azure", "GCP", "cloud security"],
        "soc": ["SOC", "incident response", "SIEM", "threat detection", "security operations", "security monitoring"],
        "penetration": ["penetration testing", "vulnerability assessment", "ethical hacking", "security testing"],
        "data": ["data", "data analysis", "data science", "analytics", "business intelligence"],
        "network": ["network", "network security", "firewall", "IDS", "IPS"],
    }

    focus_areas = []
    for i, role in enumerate(roles):
        role_lower = role.lower()
        weight = max(3 - i, 1)  # First=3, second=2, rest=1

        # Collect keywords from role name
        keywords = [role]  # Always include the full role name
        for key, kw_list in ROLE_KEYWORD_MAP.items():
            if key in role_lower:
                keywords.extend(kw_list)

        # Deduplicate while preserving order
        seen = set()
        unique_kw = []
        for kw in keywords:
            kw_lower = kw.lower()
            if kw_lower not in seen:
                seen.add(kw_lower)
                unique_kw.append(kw)

        focus_areas.append({
            "name": role,
            "weight": weight,
            "keywords": unique_kw,
        })

    # Add cert-derived keywords to the most relevant focus area
    certs = data.get("discovery", {}).get("certs", "")
    if certs:
        cert_keywords = [c.strip() for c in re.split(r"[,;\n]", certs) if c.strip()]
        if focus_areas and cert_keywords:
            # Add certs to first focus area (highest priority)
            existing = {kw.lower() for kw in focus_areas[0]["keywords"]}
            for ck in cert_keywords:
                if ck.lower() not in existing:
                    focus_areas[0]["keywords"].append(ck)
                    existing.add(ck.lower())

    # Default anti-signals
    anti_signals = [
        "sales", "presales", "pre-sales", "account executive",
        "business development", "marketing", "recruiter", "staffing",
    ]

    return {
        "min_score": 0.3,
        "focus_areas": focus_areas,
        "anti_signals": anti_signals,
    }


def build_customer_config(data: dict, slug: str) -> dict:
    """Generate a customer-specific config.yaml."""
    # Load the master config as template for API keys
    master = {}
    if TEMPLATE_CONFIG.exists():
        with open(TEMPLATE_CONFIG, "r", encoding="utf-8") as f:
            master = yaml.safe_load(f) or {}

    work_pref = data.get("workArrangement", "any")
    remote_only = work_pref == "remote"

    # Parse salary
    min_salary = 0
    sal_str = data.get("minSalary", "")
    if sal_str:
        nums = re.findall(r"[\d,]+", sal_str.replace(",", ""))
        if nums:
            min_salary = int(nums[0])

    config = {
        "candidate": {
            "name": f"{data['firstName']} {data['lastName']}",
            "location": data.get("location", ""),
            "email": data.get("email", ""),
            "base_resume": "base_resume.md",
            "resume_library": "resumes/",
            "base_cover_letter": "base_cover_letter.md",
            "cover_letter_library": "cover_letters/",
        },
        "search": {
            "queries": data.get("roles") or ["Scrum Master", "Product Manager"],
            "locations": data.get("locations") or [data.get("location", "Canada") or "Canada"],
            "remote_only": remote_only,
            "date_posted": "week",
            "results_per_query": 10,
        },
        "api_keys": {
            "jsearch_rapidapi": master.get("api_keys", {}).get("jsearch_rapidapi", ""),
            "adzuna_app_id": master.get("api_keys", {}).get("adzuna_app_id", ""),
            "adzuna_app_key": master.get("api_keys", {}).get("adzuna_app_key", ""),
        },
        "sources": master.get("sources", {
            "jsearch": {"enabled": True, "base_url": "https://jsearch.p.rapidapi.com/search"},
            "remotive": {"enabled": True, "base_url": "https://remotive.com/api/remote-jobs"},
            "adzuna": {"enabled": True, "base_url": "https://api.adzuna.com/v1/api/jobs", "country": "ca"},
            "remoteok": {"enabled": True, "base_url": "https://remoteok.com/api"},
            "themuse": {"enabled": True, "base_url": "https://www.themuse.com/api/public/jobs"},
        }),
        "output": {
            "jd_directory": "JDs",
            "state_file": "state.json",
        },
        "filters": {
            "exclude_companies": data.get("exclude", []),
            "min_salary": min_salary,
            "must_contain": [],
            "exclude_keywords": [],
        },
        "notifications": master.get("notifications", {
            "email": {
                "sender": "",
                "app_password": "",
                "recipient": data.get("email", ""),
            }
        }),
        "lifecycle": {
            "started_at": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
            "expires_at": (datetime.now(timezone.utc) + timedelta(days=20)).strftime("%Y-%m-%dT%H:%M:%SZ"),
            "daily_limit": 10,
            "total_delivered": 0,
        },
    }

    # Auto-populate relevance config from onboarding data
    relevance = _build_relevance_config(data)
    if relevance:
        config["relevance"] = relevance

    return config


def build_discovery_notes(data: dict) -> str:
    """Generate discovery notes markdown from onboarding data."""
    d = data.get("discovery", {})
    name = f"{data['firstName']} {data['lastName']}"

    md = f"# Discovery Notes: {name}\n\n"
    md += f"**Date:** {data.get('submittedAt', datetime.now().isoformat())[:10]}\n\n---\n\n"

    fields = [
        ("Team Leadership Scale", d.get("teamSize")),
        ("Largest Project/Budget", d.get("budget")),
        ("Measurable Improvements", d.get("metrics")),
        ("Certifications", d.get("certs")),
        ("Challenge Overcome", d.get("challenge")),
        ("Tools & Technologies", d.get("tools")),
        ("Industries Worked In", d.get("industries")),
        ("Underrepresented Experience", d.get("hidden")),
    ]

    for label, value in fields:
        if value:
            md += f"## {label}\n\n{value}\n\n"

    notes = data.get("prefNotes", "")
    if notes:
        md += f"## Additional Notes\n\n{notes}\n\n"

    return md


# ---------------------------------------------------------------------------
# List customers
# ---------------------------------------------------------------------------
def cmd_list(args):
    """List all customers and their status."""
    if not CUSTOMERS_DIR.exists():
        print("  No customers yet. Import one with: python manage.py import <file.json>")
        return

    print(f"\n  {'CUSTOMER':<25} {'ROLES':<30} {'JDs':<6} {'SENT':<6} {'STATUS'}")
    print(f"  {'─'*25} {'─'*30} {'─'*6} {'─'*6} {'─'*18}")

    now = datetime.now(timezone.utc)
    for cdir in sorted(CUSTOMERS_DIR.iterdir()):
        if not cdir.is_dir():
            continue

        config_path = cdir / "config.yaml"
        if not config_path.exists():
            continue

        with open(config_path, "r", encoding="utf-8") as f:
            config = yaml.safe_load(f)

        name = config.get("candidate", {}).get("name", cdir.name)
        roles = ", ".join(config.get("search", {}).get("queries", []))[:28]

        # Count JDs
        jd_count = len(list((cdir / "JDs").glob("*.md"))) if (cdir / "JDs").exists() else 0

        # Lifecycle info
        lifecycle = config.get("lifecycle", {})
        total_delivered = lifecycle.get("total_delivered", 0)
        expires_at_str = lifecycle.get("expires_at", "")

        if expires_at_str:
            expires_at = datetime.strptime(expires_at_str, "%Y-%m-%dT%H:%M:%SZ").replace(tzinfo=timezone.utc)
            days_left = (expires_at - now).days
            if days_left < 0:
                status = "EXPIRED"
            else:
                status = f"ACTIVE ({days_left}d left)"
        else:
            status = "NO LIFECYCLE"

        print(f"  {name:<25} {roles:<30} {jd_count:<6} {total_delivered:<6} {status}")

    print()


# ---------------------------------------------------------------------------
# Customer status
# ---------------------------------------------------------------------------
def cmd_status(args):
    """Show detailed status for a customer."""
    slug = args.customer
    cdir = CUSTOMERS_DIR / slug

    if not cdir.exists():
        print(f"ERROR: Customer '{slug}' not found.")
        sys.exit(1)

    config_path = cdir / "config.yaml"
    with open(config_path, "r", encoding="utf-8") as f:
        config = yaml.safe_load(f)

    name = config.get("candidate", {}).get("name", slug)

    print(f"\n  Customer: {name}")
    print(f"  Directory: {cdir}")
    print(f"  Email: {config.get('candidate', {}).get('email', 'N/A')}")
    print(f"  Roles: {', '.join(config.get('search', {}).get('queries', []))}")
    print(f"  Locations: {', '.join(config.get('search', {}).get('locations', []))}")

    jd_count = len(list((cdir / "JDs").glob("*.md"))) if (cdir / "JDs").exists() else 0
    print(f"\n  JDs fetched: {jd_count}")

    tailor_state_path = cdir / "tailor_state.json"
    if tailor_state_path.exists():
        ts = json.loads(tailor_state_path.read_text(encoding="utf-8"))
        tailored = ts.get("tailored", {})
        print(f"  Tailored: {len(tailored)}")
    else:
        print(f"  Tailored: 0")

    output_dir = cdir / "output"
    if output_dir.exists():
        packages = [d for d in output_dir.iterdir() if d.is_dir()]
        print(f"  Output packages: {len(packages)}")
        for pkg in sorted(packages)[:10]:
            print(f"    - {pkg.name}")
        if len(packages) > 10:
            print(f"    ... and {len(packages) - 10} more")

    # Lifecycle
    lifecycle = config.get("lifecycle", {})
    if lifecycle:
        now = datetime.now(timezone.utc)
        started = lifecycle.get("started_at", "N/A")
        expires_str = lifecycle.get("expires_at", "")
        daily_limit = lifecycle.get("daily_limit", "N/A")
        total_delivered = lifecycle.get("total_delivered", 0)

        print(f"\n  Lifecycle:")
        print(f"    Started: {started[:10] if started != 'N/A' else 'N/A'}")
        if expires_str:
            expires_at = datetime.strptime(expires_str, "%Y-%m-%dT%H:%M:%SZ").replace(tzinfo=timezone.utc)
            days_left = (expires_at - now).days
            if days_left < 0:
                print(f"    Expires: {expires_str[:10]} (EXPIRED)")
            else:
                print(f"    Expires: {expires_str[:10]} ({days_left} days remaining)")
        print(f"    Daily limit: {daily_limit}")
        print(f"    Total delivered: {total_delivered}")

    print()


# ---------------------------------------------------------------------------
# Run pipeline for a customer
# ---------------------------------------------------------------------------
def cmd_run(args):
    """Run the pipeline for a specific customer."""
    slug = args.customer
    cdir = CUSTOMERS_DIR / slug

    if not cdir.exists():
        print(f"ERROR: Customer '{slug}' not found.")
        sys.exit(1)

    print(f"\n  Running pipeline for: {slug}")
    print(f"  Directory: {cdir}\n")

    # Scripts check PIPELINE_WORKDIR env var to override ROOT
    env = {**os.environ, "PIPELINE_WORKDIR": str(cdir)}

    # Run the pipeline
    no_url_check = getattr(args, "no_url_check", False)
    steps = []
    if not args.skip_fetch:
        fetch_cmd = [PYTHON, str(ROOT / "fetcher.py")]
        if no_url_check:
            fetch_cmd.append("--no-url-check")
        steps.append(("FETCH", fetch_cmd))
    if not args.fetch_only:
        tailor_cmd = [PYTHON, str(ROOT / "tailor.py"), "--limit", str(args.tailor_limit)]
        if no_url_check:
            tailor_cmd.append("--no-url-check")
        steps.append(("TAILOR", tailor_cmd))
        notify_cmd = [PYTHON, str(ROOT / "notify.py")]
        if args.no_email:
            notify_cmd.append("--no-email")
        if args.notify_limit:
            notify_cmd.extend(["--limit", str(args.notify_limit)])
        if args.min_match:
            notify_cmd.extend(["--min-match", str(args.min_match)])
        if args.min_relevance:
            notify_cmd.extend(["--min-relevance", str(args.min_relevance)])
        if no_url_check:
            notify_cmd.append("--no-url-check")
        steps.append(("NOTIFY", notify_cmd))

    for name, cmd in steps:
        print(f"  [{name}]")
        result = subprocess.run(cmd, cwd=str(cdir), env=env, timeout=7200)
        if result.returncode != 0:
            print(f"  [{name}] FAILED (exit {result.returncode})")
        else:
            print(f"  [{name}] OK")
        print()


# ---------------------------------------------------------------------------
# Run all customers
# ---------------------------------------------------------------------------
def cmd_run_all(args):
    """Run the pipeline for all customers."""
    if not CUSTOMERS_DIR.exists():
        print("  No customers yet.")
        return

    customers = [d for d in sorted(CUSTOMERS_DIR.iterdir()) if d.is_dir() and (d / "config.yaml").exists()]
    print(f"\n  Running pipeline for {len(customers)} customer(s)...\n")

    now = datetime.now(timezone.utc)
    for cdir in customers:
        slug = cdir.name

        # Check lifecycle expiry
        with open(cdir / "config.yaml", "r", encoding="utf-8") as f:
            cust_config = yaml.safe_load(f)
        lifecycle = cust_config.get("lifecycle", {})
        expires_at_str = lifecycle.get("expires_at", "")
        if expires_at_str:
            expires_at = datetime.strptime(expires_at_str, "%Y-%m-%dT%H:%M:%SZ").replace(tzinfo=timezone.utc)
            if now > expires_at:
                print(f"  SKIPPING {slug} — expired on {expires_at_str[:10]}")
                print()
                continue

        print(f"{'='*60}")
        print(f"  CUSTOMER: {slug}")
        print(f"{'='*60}")

        # Reuse cmd_run logic
        args.customer = slug
        try:
            cmd_run(args)
        except Exception as e:
            print(f"  ERROR: {e}")
        print()


# ---------------------------------------------------------------------------
# Renew customer
# ---------------------------------------------------------------------------
def cmd_renew(args):
    """Extend a customer's expiry date."""
    slug = args.customer
    cdir = CUSTOMERS_DIR / slug

    if not cdir.exists():
        print(f"ERROR: Customer '{slug}' not found.")
        sys.exit(1)

    config_path = cdir / "config.yaml"
    with open(config_path, "r", encoding="utf-8") as f:
        config = yaml.safe_load(f)

    lifecycle = config.get("lifecycle", {})
    if not lifecycle:
        # Initialize lifecycle if missing
        lifecycle = {
            "started_at": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
            "expires_at": (datetime.now(timezone.utc) + timedelta(days=args.days)).strftime("%Y-%m-%dT%H:%M:%SZ"),
            "daily_limit": 10,
            "total_delivered": 0,
        }
        config["lifecycle"] = lifecycle
    else:
        # Extend from current expiry (or now if already expired)
        expires_str = lifecycle.get("expires_at", "")
        if expires_str:
            current_expiry = datetime.strptime(expires_str, "%Y-%m-%dT%H:%M:%SZ").replace(tzinfo=timezone.utc)
            base = max(current_expiry, datetime.now(timezone.utc))
        else:
            base = datetime.now(timezone.utc)
        new_expiry = base + timedelta(days=args.days)
        lifecycle["expires_at"] = new_expiry.strftime("%Y-%m-%dT%H:%M:%SZ")

    with open(config_path, "w", encoding="utf-8") as f:
        yaml.dump(config, f, default_flow_style=False, allow_unicode=True)

    print(f"\n  Renewed {slug} by {args.days} days.")
    print(f"  New expiry: {lifecycle['expires_at'][:10]}")
    print()


# ---------------------------------------------------------------------------
# Recover files
# ---------------------------------------------------------------------------
def recover_from_local(cdir: Path, file_path: Path, file_type: str = "resume"):
    """Recover a binary file from a local path into the customer's library."""
    if not file_path.exists():
        print(f"  ERROR: File not found: {file_path}")
        sys.exit(1)

    # Copy binary to customer dir
    safe_name = re.sub(r"[^a-zA-Z0-9._-]", "_", file_path.name)
    dest = cdir / safe_name
    shutil.copy2(file_path, dest)
    print(f"  Saved binary file: {safe_name} ({dest.stat().st_size} bytes)")

    # Extract text
    file_bytes = dest.read_bytes()
    file_data_b64 = base64.b64encode(file_bytes).decode("ascii")
    extracted = _extract_file_text(safe_name, file_data_b64, cdir).strip()

    if not extracted:
        print(f"  WARNING: Could not extract text from {safe_name}")
        return

    # Save to library
    if file_type == "resume":
        lib_dir = cdir / "resumes"
        lib_dir.mkdir(exist_ok=True)
        lib_file = lib_dir / "uploaded_resume.md"
        lib_file.write_text(extracted, encoding="utf-8")
        print(f"  Saved: resumes/uploaded_resume.md ({len(extracted)} chars)")
    elif file_type == "cover_letter":
        lib_dir = cdir / "cover_letters"
        lib_dir.mkdir(exist_ok=True)
        lib_file = lib_dir / "uploaded_cover_letter.md"
        lib_file.write_text(extracted, encoding="utf-8")
        print(f"  Saved: cover_letters/uploaded_cover_letter.md ({len(extracted)} chars)")


def recover_from_drive(cdir: Path, data: dict):
    """Search Google Drive for binary resume/cover letter files for this customer."""
    try:
        from auto_import import get_drive_service, find_folder
    except ImportError:
        print("  ERROR: Could not import auto_import module.")
        print("  Make sure auto_import.py is in the project root and google-api-python-client is installed.")
        return

    first = data.get("firstName", "")
    last = data.get("lastName", "")
    if not first or not last:
        print("  ERROR: onboarding.json missing firstName/lastName.")
        return

    print(f"  Searching Google Drive for files matching: *{first}_{last}*")
    service = get_drive_service()

    folders_to_search = [
        "jobRadar_Inbox", "jobRadar_Processed",
        "Onboarding_Inbox", "Onboarding_Processed",
    ]

    found_any = False
    for folder_name in folders_to_search:
        folder_id = find_folder(service, folder_name)
        if not folder_id:
            continue

        for prefix, file_type in [("resume_", "resume"), ("cover_letter_", "cover_letter")]:
            query = (
                f"'{folder_id}' in parents "
                f"and name contains '{prefix}{first}_{last}' "
                f"and trashed = false"
            )
            resp = service.files().list(
                q=query, spaces="drive",
                fields="files(id, name, mimeType, size)"
            ).execute()

            for f in resp.get("files", []):
                found_any = True
                print(f"  Found: {f['name']} in {folder_name}/ ({f.get('size', '?')} bytes)")

                # Download binary
                import io
                from googleapiclient.http import MediaIoBaseDownload
                request = service.files().get_media(fileId=f["id"])
                buffer = io.BytesIO()
                downloader = MediaIoBaseDownload(buffer, request)
                done = False
                while not done:
                    _, done = downloader.next_chunk()
                buffer.seek(0)
                file_bytes = buffer.read()

                # Save binary
                safe_name = re.sub(r"[^a-zA-Z0-9._-]", "_", f["name"])
                (cdir / safe_name).write_bytes(file_bytes)
                print(f"  Downloaded: {safe_name} ({len(file_bytes)} bytes)")

                # Extract text and save to library
                file_data_b64 = base64.b64encode(file_bytes).decode("ascii")
                extracted = _extract_file_text(safe_name, file_data_b64, cdir).strip()
                if extracted:
                    if file_type == "resume":
                        lib_dir = cdir / "resumes"
                        lib_dir.mkdir(exist_ok=True)
                        (lib_dir / "uploaded_resume.md").write_text(extracted, encoding="utf-8")
                        print(f"  Extracted: resumes/uploaded_resume.md ({len(extracted)} chars)")
                    elif file_type == "cover_letter":
                        lib_dir = cdir / "cover_letters"
                        lib_dir.mkdir(exist_ok=True)
                        (lib_dir / "uploaded_cover_letter.md").write_text(extracted, encoding="utf-8")
                        print(f"  Extracted: cover_letters/uploaded_cover_letter.md ({len(extracted)} chars)")
                else:
                    print(f"  WARNING: Could not extract text from {safe_name}")

    # Also check for full onboarding JSON with base64 data
    for folder_name in folders_to_search:
        folder_id = find_folder(service, folder_name)
        if not folder_id:
            continue
        query = (
            f"'{folder_id}' in parents "
            f"and name contains 'onboarding_{first}_{last}' "
            f"and mimeType = 'application/json' "
            f"and trashed = false"
        )
        resp = service.files().list(
            q=query, spaces="drive",
            fields="files(id, name)"
        ).execute()
        for f in resp.get("files", []):
            print(f"  Found JSON: {f['name']} in {folder_name}/")
            import io
            from googleapiclient.http import MediaIoBaseDownload
            request = service.files().get_media(fileId=f["id"])
            buffer = io.BytesIO()
            downloader = MediaIoBaseDownload(buffer, request)
            done = False
            while not done:
                _, done = downloader.next_chunk()
            buffer.seek(0)
            full_data = json.loads(buffer.read().decode("utf-8"))

            # Check if it has base64 resume data
            if full_data.get("resumeFileData"):
                found_any = True
                fname = full_data.get("resumeFileName", "") or full_data.get("resumeFile", "")
                print(f"  JSON contains resumeFileData for: {fname}")
                extracted = _extract_file_text(
                    fname, full_data["resumeFileData"], cdir
                ).strip()
                if extracted:
                    lib_dir = cdir / "resumes"
                    lib_dir.mkdir(exist_ok=True)
                    (lib_dir / "uploaded_resume.md").write_text(extracted, encoding="utf-8")
                    print(f"  Extracted: resumes/uploaded_resume.md ({len(extracted)} chars)")

            if full_data.get("coverLetterFileData"):
                found_any = True
                fname = full_data.get("coverLetterFileName", "") or full_data.get("coverLetterFile", "")
                print(f"  JSON contains coverLetterFileData for: {fname}")
                extracted = _extract_file_text(
                    fname, full_data["coverLetterFileData"], cdir
                ).strip()
                if extracted:
                    lib_dir = cdir / "cover_letters"
                    lib_dir.mkdir(exist_ok=True)
                    (lib_dir / "uploaded_cover_letter.md").write_text(extracted, encoding="utf-8")
                    print(f"  Extracted: cover_letters/uploaded_cover_letter.md ({len(extracted)} chars)")

    if not found_any:
        print(f"  No binary files found on Google Drive for {first} {last}.")
        print(f"  The file must be obtained from the customer directly.")
        print(f"  Once you have the file, run:")
        print(f"    python manage.py recover-files {cdir.name} --file /path/to/file.docx")


def cmd_recover(args):
    """Recover binary resume/cover letter files for a customer."""
    slug = args.customer
    cdir = CUSTOMERS_DIR / slug
    if not cdir.exists():
        print(f"  ERROR: Customer '{slug}' not found.")
        sys.exit(1)

    onboarding_path = cdir / "onboarding.json"
    if not onboarding_path.exists():
        print(f"  ERROR: No onboarding.json found for {slug}.")
        sys.exit(1)
    data = json.loads(onboarding_path.read_text(encoding="utf-8"))

    print(f"\n  Recovering files for: {data.get('firstName', '')} {data.get('lastName', '')}")
    print(f"  Customer directory: {cdir}\n")

    if args.file:
        recover_from_local(cdir, Path(args.file), args.type)
    else:
        recover_from_drive(cdir, data)

    print()


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(description="Customer management for job pipeline")
    subparsers = parser.add_subparsers(dest="command", help="Available commands")

    # import
    p_import = subparsers.add_parser("import", help="Import customer from onboarding JSON")
    p_import.add_argument("file", help="Path to onboarding JSON file")
    p_import.add_argument("--force", action="store_true", help="Overwrite existing customer")

    # list
    subparsers.add_parser("list", help="List all customers")

    # status
    p_status = subparsers.add_parser("status", help="Show customer status")
    p_status.add_argument("customer", help="Customer slug (e.g. john-doe)")

    # run
    p_run = subparsers.add_parser("run", help="Run pipeline for a customer")
    p_run.add_argument("customer", help="Customer slug")
    p_run.add_argument("--fetch-only", action="store_true")
    p_run.add_argument("--skip-fetch", action="store_true")
    p_run.add_argument("--tailor-limit", type=int, default=10)
    p_run.add_argument("--notify-limit", type=int, default=0, help="Max packages to include in digest (0 = unlimited)")
    p_run.add_argument("--min-match", type=int, default=0, help="Minimum match %% to include in digest (e.g. 80)")
    p_run.add_argument("--min-relevance", type=int, default=0, help="Minimum relevance %% to include in digest (e.g. 60)")
    p_run.add_argument("--no-email", action="store_true")
    p_run.add_argument("--no-url-check", action="store_true", help="Skip URL liveness checks in all phases")

    # run-all
    p_all = subparsers.add_parser("run-all", help="Run pipeline for all customers")
    p_all.add_argument("--fetch-only", action="store_true")
    p_all.add_argument("--skip-fetch", action="store_true")
    p_all.add_argument("--tailor-limit", type=int, default=10)
    p_all.add_argument("--notify-limit", type=int, default=0, help="Max packages to include in digest (0 = unlimited)")
    p_all.add_argument("--min-match", type=int, default=0, help="Minimum match %% to include in digest (e.g. 80)")
    p_all.add_argument("--min-relevance", type=int, default=0, help="Minimum relevance %% to include in digest (e.g. 60)")
    p_all.add_argument("--no-email", action="store_true")
    p_all.add_argument("--no-url-check", action="store_true", help="Skip URL liveness checks in all phases")

    # renew
    p_renew = subparsers.add_parser("renew", help="Extend customer expiry date")
    p_renew.add_argument("customer", help="Customer slug")
    p_renew.add_argument("--days", type=int, default=20, help="Days to extend (default: 20)")

    # recover-files
    p_recover = subparsers.add_parser("recover-files", help="Recover binary files for a customer")
    p_recover.add_argument("customer", help="Customer slug")
    p_recover.add_argument("--file", type=str, help="Local file path to recover from")
    p_recover.add_argument("--type", choices=["resume", "cover_letter"], default="resume",
                           help="File type (default: resume)")

    args = parser.parse_args()

    if args.command == "import":
        cmd_import(args)
    elif args.command == "list":
        cmd_list(args)
    elif args.command == "status":
        cmd_status(args)
    elif args.command == "run":
        cmd_run(args)
    elif args.command == "run-all":
        cmd_run_all(args)
    elif args.command == "renew":
        cmd_renew(args)
    elif args.command == "recover-files":
        cmd_recover(args)
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
